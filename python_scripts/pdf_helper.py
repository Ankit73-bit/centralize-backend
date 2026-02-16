#!/usr/bin/env python3
"""
DOCX to PDF Converter with Formatting Preservation
Uses python-docx and reportlab to maintain document structure
"""
import sys
import json
import os
from io import BytesIO

try:
    from docx import Document
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError as e:
    print(json.dumps({
        'success': False,
        'error': f'Missing required library: {str(e)}',
        'installed': 'python-docx, reportlab should be installed'
    }))
    sys.exit(1)


def parse_docx_paragraph(para):
    """Parse paragraph and extract text with formatting"""
    text_parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
            
        # Build HTML-like formatting
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.underline:
            text = f"<u>{text}</u>"
            
        text_parts.append(text)
    
    return ''.join(text_parts)


def get_alignment(para):
    """Get paragraph alignment"""
    if para.alignment is None:
        return TA_LEFT
    
    alignment_map = {
        0: TA_LEFT,      # LEFT
        1: TA_CENTER,    # CENTER
        2: TA_RIGHT,     # RIGHT
        3: TA_JUSTIFY,   # JUSTIFY
    }
    
    # WD_ALIGN_PARAGRAPH values
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return TA_CENTER
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return TA_RIGHT
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return TA_JUSTIFY
    except:
        pass
    
    return TA_LEFT


def docx_to_pdf(input_path, output_path):
    """
    Convert DOCX to PDF preserving formatting
    """
    try:
        # Load DOCX document
        doc = Document(input_path)
        
        # Create PDF
        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        story = []
        
        # Process each paragraph
        for para in doc.paragraphs:
            text = parse_docx_paragraph(para)
            
            if not text.strip():
                # Empty paragraph - add small space
                story.append(Spacer(1, 0.2 * inch))
                continue
            
            # Determine style based on paragraph style
            style_name = para.style.name
            
            if 'Heading 1' in style_name or 'Title' in style_name:
                style = ParagraphStyle(
                    'CustomHeading1',
                    parent=styles['Heading1'],
                    fontSize=18,
                    textColor=colors.black,
                    spaceAfter=12,
                    spaceBefore=12,
                    alignment=get_alignment(para)
                )
            elif 'Heading 2' in style_name:
                style = ParagraphStyle(
                    'CustomHeading2',
                    parent=styles['Heading2'],
                    fontSize=14,
                    textColor=colors.black,
                    spaceAfter=10,
                    spaceBefore=10,
                    alignment=get_alignment(para)
                )
            elif 'Heading 3' in style_name:
                style = ParagraphStyle(
                    'CustomHeading3',
                    parent=styles['Heading3'],
                    fontSize=12,
                    textColor=colors.black,
                    spaceAfter=8,
                    spaceBefore=8,
                    alignment=get_alignment(para)
                )
            else:
                # Normal paragraph
                style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.black,
                    spaceAfter=6,
                    alignment=get_alignment(para)
                )
            
            # Create paragraph
            try:
                p = Paragraph(text, style)
                story.append(p)
            except Exception as e:
                # Fallback to plain text if formatting fails
                p = Paragraph(text.replace('<b>', '').replace('</b>', '')
                             .replace('<i>', '').replace('</i>', '')
                             .replace('<u>', '').replace('</u>', ''), style)
                story.append(p)
        
        # Process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:
                # Create PDF table
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
                story.append(Spacer(1, 0.3 * inch))
        
        # Build PDF
        pdf_doc.build(story)
        
        # Get file size
        output_size = os.path.getsize(output_path)
        
        return {
            'success': True,
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'output_size': output_size,
            'format': 'pdf'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def main():
    """Main entry point"""
    if len(sys.argv) < 3:
        print(json.dumps({
            'success': False,
            'error': 'Usage: python docx_to_pdf.py <input.docx> <output.pdf>'
        }))
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not os.path.exists(input_path):
        print(json.dumps({
            'success': False,
            'error': f'Input file not found: {input_path}'
        }))
        sys.exit(1)
    
    result = docx_to_pdf(input_path, output_path)
    print(json.dumps(result))
    
    sys.exit(0 if result['success'] else 1)


if __name__ == '__main__':
    main()